"""File converter module for various document formats to text"""

import logging
from pathlib import Path
from typing import Optional, List
from datetime import datetime

logger = logging.getLogger(__name__)


class FileConverter:
    """다양한 파일 형식을 텍스트로 변환하는 클래스"""
    
    # 지원하는 파일 확장자
    SUPPORTED_EXTENSIONS = {
        'text': ['.txt', '.log', '.md', '.csv'],
        'pdf': ['.pdf'],
        'word': ['.doc', '.docx'],
        'hwp': ['.hwp'],
        'excel': ['.xls', '.xlsx'],
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
    }
    
    def __init__(self):
        """Initialize file converter with optional dependency checks"""
        self._check_dependencies()
        # Initialize system converter for fallback conversions
        try:
            from .system_file_converter import SystemFileConverter
            self.system_converter = SystemFileConverter()
        except ImportError:
            logger.warning("SystemFileConverter not available")
            self.system_converter = None
    
    def _check_dependencies(self):
        """Check for required dependencies"""
        self.dependencies = {
            'pdf': False,
            'word': False,
            'hwp': False,
            'excel': False,
            'ocr': False
        }
        
        try:
            import pypdf
            self.dependencies['pdf'] = True
        except ImportError:
            try:
                import PyPDF2
                self.dependencies['pdf'] = True
            except ImportError:
                logger.warning("pypdf/PyPDF2 not installed. PDF conversion unavailable.")
        
        try:
            import docx
            self.dependencies['word'] = True
        except ImportError:
            logger.warning("python-docx not installed. Word conversion unavailable.")
        
        try:
            import pyhwp
            self.dependencies['hwp'] = True
        except ImportError:
            logger.warning("pyhwp not installed. HWP conversion unavailable.")
        
        try:
            import openpyxl
            self.dependencies['excel'] = True
        except ImportError:
            logger.warning("openpyxl not installed. Excel conversion unavailable.")
        
        try:
            import pytesseract
            from PIL import Image
            self.dependencies['ocr'] = True
        except ImportError:
            logger.warning("pytesseract/PIL not installed. OCR unavailable.")
    
    def convert_to_text(self, file_path: Path) -> str:
        """
        Convert file to text based on its extension
        
        Args:
            file_path: Path to file to convert
            
        Returns:
            Extracted text content
        """
        if not file_path.exists():
            return f"Error: File not found - {file_path}"
        
        file_ext = file_path.suffix.lower()
        
        try:
            # 텍스트 파일
            if file_ext in self.SUPPORTED_EXTENSIONS['text']:
                return self._convert_text_file(file_path)
            
            # PDF
            elif file_ext in self.SUPPORTED_EXTENSIONS['pdf']:
                if self.dependencies['pdf']:
                    return self._convert_pdf_to_text(file_path)
                else:
                    # Fallback to system converter
                    from .system_file_converter import SystemFileConverter
                    system_converter = SystemFileConverter()
                    return system_converter.convert_pdf_to_text(file_path)
            
            # Word
            elif file_ext in self.SUPPORTED_EXTENSIONS['word']:
                if self.dependencies['word']:
                    return self._convert_word_to_text(file_path)
                else:
                    # Fallback to system converter
                    from .system_file_converter import SystemFileConverter
                    system_converter = SystemFileConverter()
                    if file_ext == '.docx':
                        return system_converter.convert_docx_to_text(file_path)
                    else:
                        return system_converter.convert_doc_to_text(file_path)
            
            # HWP (한글)
            elif file_ext in self.SUPPORTED_EXTENSIONS['hwp']:
                if self.dependencies['hwp']:
                    return self._convert_hwp_to_text(file_path)
                else:
                    # Fallback to system converter
                    from .system_file_converter import SystemFileConverter
                    system_converter = SystemFileConverter()
                    return system_converter.convert_hwp_to_text(file_path)
            
            # Excel
            elif file_ext in self.SUPPORTED_EXTENSIONS['excel']:
                if self.dependencies['excel']:
                    return self._convert_excel_to_text(file_path)
                else:
                    return "Excel conversion unavailable. Install openpyxl."
            
            # 이미지
            elif file_ext in self.SUPPORTED_EXTENSIONS['image']:
                if self.dependencies['ocr']:
                    return self._convert_image_to_text(file_path)
                else:
                    return "OCR unavailable. Install pytesseract and PIL."
            
            else:
                return f"Unsupported file format: {file_ext}"
                
        except Exception as e:
            logger.error(f"Failed to convert {file_path}: {str(e)}")
            return f"Conversion error: {str(e)}"
    
    def _convert_text_file(self, file_path: Path) -> str:
        """Convert text file with encoding detection"""
        encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    logger.info(f"Successfully read {file_path} with {encoding} encoding")
                    return content
            except UnicodeDecodeError:
                continue
        
        return "Unable to decode text file with supported encodings"
    
    def _convert_pdf_to_text(self, file_path: Path) -> str:
        """Convert PDF to text using pypdf or PyPDF2"""
        try:
            import pypdf as pdf_lib
        except ImportError:
            import PyPDF2 as pdf_lib
        
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pdf_lib.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---")
                        text_parts.append(text)
                
                if not text_parts:
                    return "No text content found in PDF"
                
                return '\n\n'.join(text_parts)
                
        except Exception as e:
            logger.error(f"PDF conversion error: {str(e)}")
            return f"PDF conversion failed: {str(e)}"
    
    def _convert_word_to_text(self, file_path: Path) -> str:
        """Convert Word document to text using python-docx"""
        import docx
        
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # 문단 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 표 추출
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append("\n[Table Content]")
                    text_parts.append(table_text)
            
            if not text_parts:
                return "No text content found in Word document"
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Word conversion error: {str(e)}")
            return f"Word conversion failed: {str(e)}"
    
    def _extract_table_text(self, table) -> str:
        """Extract text from Word table"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            if any(cells):  # 빈 행 제외
                rows.append(' | '.join(cells))
        return '\n'.join(rows)
    
    def _convert_hwp_to_text(self, file_path: Path) -> str:
        """Convert HWP (한글) document to text"""
        try:
            # Method 1: Try system converter first (hwp5html for better table support)
            if self.system_converter:
                try:
                    result = self.system_converter.convert_hwp_to_text(file_path)
                    if result and not result.startswith("Unable to convert"):
                        return result
                except Exception as e:
                    logger.warning(f"System converter failed: {str(e)}")
            
            # Method 2: Try pyhwp (if available)
            try:
                import pyhwp
                
                with pyhwp.HWPReader(str(file_path)) as reader:
                    text_parts = []
                    for paragraph in reader.paragraphs:
                        text_parts.append(paragraph.text)
                    
                    if text_parts:
                        return '\n\n'.join(text_parts)
                    
            except Exception as e:
                logger.warning(f"pyhwp failed: {str(e)}")
            
            # Method 3: Try hwp5 (alternative library)
            try:
                import hwp5
                
                hwp = hwp5.HWP5File(str(file_path))
                text_parts = []
                
                for section in hwp.bodytext.sections:
                    for paragraph in section:
                        text = ''.join(record.text for record in paragraph 
                                     if hasattr(record, 'text'))
                        if text.strip():
                            text_parts.append(text)
                
                if text_parts:
                    return '\n\n'.join(text_parts)
                    
            except Exception as e:
                logger.warning(f"hwp5 failed: {str(e)}")
            
            # Method 3: Basic binary extraction (fallback)
            return self._extract_hwp_text_basic(file_path)
            
        except Exception as e:
            logger.error(f"HWP conversion error: {str(e)}")
            return f"HWP conversion failed: {str(e)}"
    
    def _extract_hwp_text_basic(self, file_path: Path) -> str:
        """Basic HWP text extraction (fallback method)"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                
            # HWP 파일은 압축된 형식이므로 간단한 텍스트 추출
            text_parts = []
            
            # UTF-16 디코딩 시도
            try:
                text = content.decode('utf-16-le', errors='ignore')
                # 인쇄 가능한 문자만 추출
                printable_text = ''.join(c for c in text if c.isprintable() or c.isspace())
                if printable_text.strip():
                    text_parts.append(printable_text)
            except:
                pass
            
            # UTF-8 디코딩 시도
            try:
                text = content.decode('utf-8', errors='ignore')
                printable_text = ''.join(c for c in text if c.isprintable() or c.isspace())
                if printable_text.strip():
                    text_parts.append(printable_text)
            except:
                pass
            
            if text_parts:
                return "Note: Basic text extraction (may contain artifacts)\n\n" + '\n'.join(text_parts)
            else:
                return "Unable to extract text from HWP file. Install pyhwp or hwp5 for better results."
                
        except Exception as e:
            return f"Basic HWP extraction failed: {str(e)}"
    
    def _convert_excel_to_text(self, file_path: Path) -> str:
        """Convert Excel file to text using openpyxl"""
        import openpyxl
        
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"=== Sheet: {sheet_name} ==="]
                
                # 데이터가 있는 영역만 처리
                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    
                    if row_data:  # 빈 행 제외
                        sheet_text.append('\t'.join(row_data))
                
                if len(sheet_text) > 1:  # 헤더만 있는 경우 제외
                    text_parts.append('\n'.join(sheet_text))
            
            if not text_parts:
                return "No data found in Excel file"
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"Excel conversion error: {str(e)}")
            return f"Excel conversion failed: {str(e)}"
    
    def _convert_image_to_text(self, file_path: Path) -> str:
        """Convert image to text using OCR"""
        try:
            import pytesseract
            from PIL import Image
            
            # 이미지 열기
            image = Image.open(file_path)
            
            # 이미지 전처리 (선택사항)
            # 그레이스케일 변환
            if image.mode != 'L':
                image = image.convert('L')
            
            # OCR 실행 (한국어 + 영어)
            try:
                text = pytesseract.image_to_string(image, lang='kor+eng')
            except:
                # 한국어 데이터가 없으면 영어만
                text = pytesseract.image_to_string(image, lang='eng')
            
            if text.strip():
                return text
            else:
                return "No text found in image"
                
        except Exception as e:
            logger.error(f"OCR error: {str(e)}")
            return f"OCR failed: {str(e)}"
    
    def save_as_text(
        self,
        file_path: Path,
        text_content: str,
        original_filename: Optional[str] = None
    ) -> Path:
        """
        Save converted text to file
        
        Args:
            file_path: Original file path
            text_content: Converted text content
            original_filename: Optional original filename for metadata
            
        Returns:
            Path to saved text file
        """
        # 텍스트 파일 경로 생성
        text_file_path = file_path.with_suffix('.txt')
        
        # 중복 방지
        counter = 1
        while text_file_path.exists():
            stem = file_path.stem
            text_file_path = file_path.parent / f"{stem}_{counter}.txt"
            counter += 1
        
        # 메타데이터 포함하여 저장
        try:
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(f"Original file: {original_filename or file_path.name}\n")
                f.write(f"Converted at: {datetime.now()}\n")
                f.write(f"Source path: {file_path}\n")
                f.write("=" * 80 + "\n\n")
                f.write(text_content)
            
            logger.info(f"Saved text to {text_file_path}")
            return text_file_path
            
        except Exception as e:
            logger.error(f"Failed to save text file: {str(e)}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Get list of all supported file extensions"""
        formats = []
        for ext_list in self.SUPPORTED_EXTENSIONS.values():
            formats.extend(ext_list)
        return sorted(list(set(formats)))
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file format is supported"""
        return file_path.suffix.lower() in self.get_supported_formats()