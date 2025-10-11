"""Python library-based file converter"""

import logging
from pathlib import Path
from typing import Optional, Dict, List
from .base import BaseConverter

logger = logging.getLogger(__name__)


class PythonLibraryConverter(BaseConverter):
    """File converter using Python libraries"""

    # Supported file extensions by category
    SUPPORTED_EXTENSIONS = {
        'text': ['.txt', '.log', '.md', '.csv'],
        'pdf': ['.pdf'],
        'word': ['.doc', '.docx'],
        'hwp': ['.hwp'],
        'excel': ['.xls', '.xlsx'],
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
    }

    def __init__(self):
        """Initialize converter and check dependencies"""
        self.dependencies = self._check_dependencies()

    def _check_dependencies(self) -> Dict[str, bool]:
        """Check for available Python libraries"""
        deps = {
            'pdf': False,
            'word': False,
            'hwp': False,
            'excel': False,
            'ocr': False
        }

        # Check PDF libraries
        try:
            import pypdf
            deps['pdf'] = True
        except ImportError:
            try:
                import PyPDF2
                deps['pdf'] = True
            except ImportError:
                logger.debug("pypdf/PyPDF2 not installed")

        # Check Word library
        try:
            import docx
            deps['word'] = True
        except ImportError:
            logger.debug("python-docx not installed")

        # Check HWP library
        try:
            import pyhwp
            deps['hwp'] = True
        except ImportError:
            logger.debug("pyhwp not installed")

        # Check Excel library
        try:
            import openpyxl
            deps['excel'] = True
        except ImportError:
            logger.debug("openpyxl not installed")

        # Check OCR libraries
        try:
            import pytesseract
            from PIL import Image
            deps['ocr'] = True
        except ImportError:
            logger.debug("pytesseract/PIL not installed")

        return deps

    def supports_format(self, file_extension: str) -> bool:
        """Check if converter supports the format"""
        ext = file_extension.lower()

        # Check each category
        for category, extensions in self.SUPPORTED_EXTENSIONS.items():
            if ext in extensions:
                # Check if we have the required dependency
                if category == 'text':
                    return True  # Always support text files
                elif category == 'pdf':
                    return self.dependencies['pdf']
                elif category == 'word':
                    return self.dependencies['word']
                elif category == 'hwp':
                    return self.dependencies['hwp']
                elif category == 'excel':
                    return self.dependencies['excel']
                elif category == 'image':
                    return self.dependencies['ocr']

        return False

    def convert_to_text(self, file_path: Path) -> Optional[str]:
        """Convert file to text"""
        if not self.validate_file(file_path):
            logger.error(f"Invalid file: {file_path}")
            return None

        file_ext = file_path.suffix.lower()

        try:
            # Text files
            if file_ext in self.SUPPORTED_EXTENSIONS['text']:
                return self._convert_text_file(file_path)

            # PDF files
            elif file_ext in self.SUPPORTED_EXTENSIONS['pdf']:
                if self.dependencies['pdf']:
                    return self._convert_pdf(file_path)

            # Word files
            elif file_ext in self.SUPPORTED_EXTENSIONS['word']:
                if self.dependencies['word']:
                    return self._convert_word(file_path)

            # HWP files
            elif file_ext in self.SUPPORTED_EXTENSIONS['hwp']:
                if self.dependencies['hwp']:
                    return self._convert_hwp(file_path)

            # Excel files
            elif file_ext in self.SUPPORTED_EXTENSIONS['excel']:
                if self.dependencies['excel']:
                    return self._convert_excel(file_path)

            # Image files
            elif file_ext in self.SUPPORTED_EXTENSIONS['image']:
                if self.dependencies['ocr']:
                    return self._convert_image_ocr(file_path)

        except Exception as e:
            logger.error(f"Error converting {file_path}: {str(e)}")
            return None

        return None

    def _convert_text_file(self, file_path: Path) -> Optional[str]:
        """Convert text file to string"""
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Fallback to cp949 for Korean files
                with open(file_path, 'r', encoding='cp949') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return None

    def _convert_pdf(self, file_path: Path) -> Optional[str]:
        """Convert PDF to text using pypdf or PyPDF2"""
        try:
            # Try pypdf first (newer version)
            try:
                import pypdf
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = pypdf.PdfReader(pdf_file)
                    text_content = []

                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text_content.append(page.extract_text())

                    return '\n'.join(text_content)

            except ImportError:
                # Fallback to PyPDF2
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text_content = []

                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text_content.append(page.extract_text())

                    return '\n'.join(text_content)

        except Exception as e:
            logger.error(f"Error converting PDF {file_path}: {str(e)}")
            return None

    def _convert_word(self, file_path: Path) -> Optional[str]:
        """Convert Word document to text"""
        try:
            import docx
            doc = docx.Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_content.append(cell.text)

            return '\n'.join(text_content)

        except Exception as e:
            logger.error(f"Error converting Word document {file_path}: {str(e)}")
            return None

    def _convert_hwp(self, file_path: Path) -> Optional[str]:
        """Convert HWP document to text"""
        try:
            import pyhwp
            hwp = pyhwp.HWP(str(file_path))
            return hwp.get_text()
        except Exception as e:
            logger.error(f"Error converting HWP {file_path}: {str(e)}")
            return None

    def _convert_excel(self, file_path: Path) -> Optional[str]:
        """Convert Excel file to text"""
        try:
            import openpyxl
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    if row_data:
                        text_content.append('\t'.join(row_data))

            workbook.close()
            return '\n'.join(text_content)

        except Exception as e:
            logger.error(f"Error converting Excel {file_path}: {str(e)}")
            return None

    def _convert_image_ocr(self, file_path: Path) -> Optional[str]:
        """Convert image to text using OCR"""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='kor+eng')
            return text

        except Exception as e:
            logger.error(f"Error performing OCR on {file_path}: {str(e)}")
            return None